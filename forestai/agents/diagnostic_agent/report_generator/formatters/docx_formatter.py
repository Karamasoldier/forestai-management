# -*- coding: utf-8 -*-
"""
Module du formateur DOCX pour les rapports de diagnostic.
"""

import logging
from typing import Dict, Any, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from forestai.agents.diagnostic_agent.report_generator.formatters.base_formatter import BaseFormatter

logger = logging.getLogger(__name__)

class DOCXFormatter(BaseFormatter):
    """Formateur de rapports au format DOCX."""
    
    def __init__(self, templates_dir: Path, output_dir: Path):
        """Initialise le formateur DOCX.
        
        Args:
            templates_dir: Répertoire des templates
            output_dir: Répertoire de sortie des rapports
        """
        super().__init__(templates_dir, output_dir)
    
    def generate(self, context: Dict[str, Any], template_name: str, output_path: Optional[Path] = None) -> Path:
        """Génère un rapport au format DOCX.
        
        Args:
            context: Contexte de données pour le template
            template_name: Nom du template à utiliser
            output_path: Chemin du fichier de sortie (optionnel)
            
        Returns:
            Chemin du fichier généré
        """
        # Vérifier si un modèle DOCX existe déjà
        template_path = self.templates_dir / template_name
        
        # Créer un nouveau document ou utiliser un modèle existant
        if template_path.exists():
            doc = Document(template_path)
        else:
            doc = self._create_new_document(context)
            
        # Déterminer le chemin de sortie si non spécifié
        if output_path is None:
            parcel_id = context.get("parcel_id", "unknown")
            timestamp = context.get("date", "").replace(" ", "_").replace(":", "").replace("/", "")
            output_path = self.output_dir / f"diagnostic_{parcel_id}_{timestamp}.docx"
        
        # Sauvegarder le document
        doc.save(str(output_path))
        
        logger.info(f"Rapport DOCX généré: {output_path}")
        return output_path
    
    def _create_new_document(self, context: Dict[str, Any]) -> Document:
        """Crée un nouveau document DOCX à partir du contexte.
        
        Args:
            context: Contexte de données pour le template
            
        Returns:
            Document DOCX créé
        """
        doc = Document()
        
        # Style du titre
        title = doc.add_heading(context.get("title", "Rapport de diagnostic forestier"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Sous-titre
        subtitle = doc.add_paragraph(context.get("subtitle", "Analyse et recommandations"))
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informations générales
        doc.add_paragraph(f"Date: {context.get('date', '')}")
        doc.add_paragraph(f"Référence: {context.get('reference', '')}")
        
        # Séparateur
        doc.add_paragraph("=" * 50)
        
        # Résumé
        doc.add_heading("Résumé du diagnostic", level=1)
        doc.add_paragraph(context.get("summary", ""))
        
        # Points clés
        if "summary_highlights" in context and context["summary_highlights"]:
            doc.add_paragraph("Points clés:")
            for highlight in context["summary_highlights"]:
                p = doc.add_paragraph()
                p.add_run("• ").bold = True
                p.add_run(highlight)
        
        # Informations sur la parcelle
        doc.add_heading("Informations sur la parcelle", level=1)
        table = doc.add_table(rows=3, cols=4)
        
        # En-têtes et valeurs
        cells = table.rows[0].cells
        cells[0].text = "Identifiant"
        cells[1].text = context.get("parcel_id", "")
        cells[2].text = "Surface"
        cells[3].text = f"{context.get('parcel_area', '')} ha"
        
        cells = table.rows[1].cells
        cells[0].text = "Commune"
        cells[1].text = context.get("parcel_commune", "")
        cells[2].text = "Propriétaire"
        cells[3].text = context.get("parcel_owner", "")
        
        cells = table.rows[2].cells
        cells[0].text = "Type de peuplement"
        cells[1].text = context.get("forest_type", "")
        cells[2].text = "Exposition"
        cells[3].text = context.get("parcel_exposition", "")
        
        # Styler la table
        table.style = 'Table Grid'
        
        # Ajouter les sections supplémentaires en fonction des données disponibles
        self._add_inventory_section(doc, context)
        self._add_climate_section(doc, context)
        self._add_recommendations_section(doc, context)
        
        # Pied de page
        doc.add_paragraph("")
        footer = doc.add_paragraph(f"Rapport généré par {context.get('company_name', 'ForestAI')} - {context.get('date', '')}")
        footer.add_run(f"\n{context.get('company_contact', '')} | {context.get('company_website', '')}")
        
        return doc
    
    def _add_inventory_section(self, doc: Document, context: Dict[str, Any]):
        """Ajoute la section d'analyse d'inventaire au document."""
        inventory_analysis = context.get("inventory_analysis", {})
        if not inventory_analysis:
            return
            
        doc.add_heading("Analyse de l'inventaire forestier", level=1)
        
        # Vue d'ensemble
        doc.add_heading("Vue d'ensemble", level=2)
        summary = inventory_analysis.get("summary", {})
        doc.add_paragraph(
            f"L'inventaire comprend un total de {summary.get('total_trees', 0)} arbres "
            f"répartis sur {summary.get('unique_species', 0)} espèces différentes."
        )
        
        # Analyse par espèce
        if "species_analysis" in inventory_analysis:
            doc.add_heading("Analyse par espèce", level=2)
            table = doc.add_table(rows=1, cols=6)
            
            # En-têtes
            header_cells = table.rows[0].cells
            header_cells[0].text = "Espèce"
            header_cells[1].text = "Nombre"
            header_cells[2].text = "Pourcentage"
            header_cells[3].text = "Diamètre moyen (cm)"
            header_cells[4].text = "Hauteur moyenne (m)"
            header_cells[5].text = "Volume moyen (m³)"
            
            # Lignes de données
            for species, data in inventory_analysis["species_analysis"].items():
                row = table.add_row().cells
                row[0].text = species
                row[1].text = str(data.get("count", 0))
                row[2].text = f"{data.get('percentage', 0):.1f}%"
                
                diameter = data.get("diameter", {})
                row[3].text = f"{diameter.get('mean', 0):.1f}" if diameter else "-"
                
                height = data.get("height", {})
                row[4].text = f"{height.get('mean', 0):.1f}" if height else "-"
                
                volume = data.get("volume", {})
                row[5].text = f"{volume.get('mean', 0):.2f}" if volume else "-"
            
            # Styler la table
            table.style = 'Table Grid'
            
        # Métriques à l'hectare
        if "per_hectare" in inventory_analysis:
            doc.add_heading("Métriques à l'hectare", level=2)
            per_hectare = inventory_analysis["per_hectare"]
            
            table = doc.add_table(rows=3, cols=2)
            
            cells = table.rows[0].cells
            cells[0].text = "Densité (arbres/ha)"
            cells[1].text = f"{per_hectare.get('density', 0):.0f}"
            
            cells = table.rows[1].cells
            cells[0].text = "Surface terrière (m²/ha)"
            cells[1].text = f"{per_hectare.get('basal_area', 0):.1f}"
            
            cells = table.rows[2].cells
            cells[0].text = "Volume sur pied (m³/ha)"
            cells[1].text = f"{per_hectare.get('volume', 0):.1f}"
            
            # Styler la table
            table.style = 'Table Grid'
    
    def _add_climate_section(self, doc: Document, context: Dict[str, Any]):
        """Ajoute la section d'analyse climatique au document."""
        climate_analysis = context.get("climate_analysis", {})
        if not climate_analysis:
            return
            
        doc.add_heading("Analyse climatique", level=1)
        
        # Climat actuel
        if "current" in climate_analysis:
            doc.add_heading("Climat actuel", level=2)
            current = climate_analysis["current"]
            
            table = doc.add_table(rows=3, cols=2)
            
            cells = table.rows[0].cells
            cells[0].text = "Température moyenne annuelle"
            cells[1].text = f"{current.get('temperature_avg', 0):.1f} °C"
            
            cells = table.rows[1].cells
            cells[0].text = "Précipitations annuelles"
            cells[1].text = f"{current.get('precipitation_annual', 0):.0f} mm"
            
            cells = table.rows[2].cells
            cells[0].text = "Indice de sécheresse estival"
            cells[1].text = f"{current.get('drought_index', 0):.1f}"
            
            # Styler la table
            table.style = 'Table Grid'
        
        # Projections climatiques
        if "future" in climate_analysis and "current" in climate_analysis:
            doc.add_heading("Projections climatiques (2050)", level=2)
            future = climate_analysis["future"]
            current = climate_analysis["current"]
            
            table = doc.add_table(rows=3, cols=3)
            
            # En-têtes
            header_cells = table.rows[0].cells
            header_cells[0].text = "Paramètre"
            header_cells[1].text = "Valeur"
            header_cells[2].text = "Variation"
            
            # Température
            cells = table.rows[0].cells
            cells[0].text = "Température moyenne annuelle"
            cells[1].text = f"{future.get('temperature_avg', 0):.1f} °C"
            temp_diff = future.get('temperature_avg', 0) - current.get('temperature_avg', 0)
            cells[2].text = f"{temp_diff:+.1f} °C"
            
            # Précipitations
            cells = table.rows[1].cells
            cells[0].text = "Précipitations annuelles"
            cells[1].text = f"{future.get('precipitation_annual', 0):.0f} mm"
            precip_diff = future.get('precipitation_annual', 0) - current.get('precipitation_annual', 0)
            cells[2].text = f"{precip_diff:+.0f} mm"
            
            # Sécheresse
            cells = table.rows[2].cells
            cells[0].text = "Indice de sécheresse estival"
            cells[1].text = f"{future.get('drought_index', 0):.1f}"
            drought_diff = future.get('drought_index', 0) - current.get('drought_index', 0)
            cells[2].text = f"{drought_diff:+.1f}"
            
            # Styler la table
            table.style = 'Table Grid'
    
    def _add_recommendations_section(self, doc: Document, context: Dict[str, Any]):
        """Ajoute la section de recommandations au document."""
        recommendations = context.get("recommendations", {})
        if not recommendations:
            return
            
        doc.add_heading("Recommandations", level=1)
        
        # Essences recommandées
        if "species" in recommendations:
            doc.add_heading("Essences recommandées", level=2)
            
            table = doc.add_table(rows=1, cols=4)
            
            # En-têtes
            header_cells = table.rows[0].cells
            header_cells[0].text = "Essence"
            header_cells[1].text = "Adaptation climatique"
            header_cells[2].text = "Potentiel productif"
            header_cells[3].text = "Commentaires"
            
            # Données
            for sp in recommendations["species"]:
                row = table.add_row().cells
                row[0].text = sp.get("name", "")
                row[1].text = f"{sp.get('climate_score', 0)}/5"
                row[2].text = f"{sp.get('productivity_score', 0)}/5"
                row[3].text = sp.get("comments", "")
            
            # Styler la table
            table.style = 'Table Grid'
        
        # Interventions sylvicoles
        if "operations" in recommendations:
            doc.add_heading("Interventions sylvicoles proposées", level=2)
            
            table = doc.add_table(rows=1, cols=4)
            
            # En-têtes
            header_cells = table.rows[0].cells
            header_cells[0].text = "Opération"
            header_cells[1].text = "Priorité"
            header_cells[2].text = "Période"
            header_cells[3].text = "Description"
            
            # Données
            for op in recommendations["operations"]:
                row = table.add_row().cells
                row[0].text = op.get("name", "")
                row[1].text = op.get("priority", "")
                row[2].text = op.get("timeframe", "")
                row[3].text = op.get("description", "")
            
            # Styler la table
            table.style = 'Table Grid'
