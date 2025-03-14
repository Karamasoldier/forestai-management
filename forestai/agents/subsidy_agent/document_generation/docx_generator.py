"""
Module pour la génération de documents DOCX liés aux subventions forestières.

Ce module permet de générer des documents Word (DOCX) pour les demandes
de subvention, les rapports d'éligibilité et les synthèses des subventions disponibles.
"""

import os
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from forestai.core.utils.logging_config import LoggingConfig


class DOCXGenerator:
    """
    Générateur de documents DOCX pour les subventions forestières.
    """
    
    def __init__(self, config: Dict[str, Any]):
        """
        Initialise le générateur de documents DOCX.
        
        Args:
            config: Dictionnaire de configuration avec les paramètres suivants:
                - templates_dir: Répertoire contenant les templates DOCX
                - output_dir: Répertoire où enregistrer les documents générés
                - logo_path: Chemin vers le logo à inclure dans les documents
        """
        self.config = config
        self.templates_dir = config.get("templates_dir", "forestai/agents/subsidy_agent/templates")
        self.output_dir = config.get("output_dir", "data/outputs/subsidies")
        self.logo_path = config.get("logo_path", "forestai/agents/subsidy_agent/templates/logo.png")
        
        # Configuration du logging
        self.logger = LoggingConfig.get_instance().get_logger(
            self.__class__.__name__,
            module=__name__
        )
        self.logger.info("Initialisation du générateur DOCX")
    
    def generate_application(self, data: Dict[str, Any], output_path: str) -> str:
        """
        Génère un dossier de demande au format DOCX.
        
        Args:
            data: Données pour le document
            output_path: Chemin où enregistrer le document
        
        Returns:
            Chemin du fichier généré
        """
        self.logger.info("Génération d'un dossier de demande DOCX")
        
        # Vérifier si un template existe
        template_path = os.path.join(self.templates_dir, "application_template.docx")
        if os.path.exists(template_path):
            doc = Document(template_path)
            self.logger.info(f"Utilisation du template: {template_path}")
        else:
            doc = Document()
            self.logger.warning("Template non trouvé, création d'un nouveau document")
        
        # Extraire les données nécessaires
        subsidy = data["subsidy"]
        parcel = data["parcel"]
        project = data["project"]
        applicant = data["applicant"]
        generated_at = data.get("generated_at", datetime.now().strftime("%d/%m/%Y"))
        
        # Formatage du document
        self._format_document_styles(doc)
        
        # Titre du document
        title = doc.add_heading("Dossier de demande de subvention", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(subsidy["title"], level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date de génération
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(f"Document généré le {generated_at}")
        
        doc.add_paragraph()  # Espace
        
        # Informations sur le demandeur
        doc.add_heading("Informations sur le demandeur", level=2)
        
        applicant_table = doc.add_table(rows=5, cols=2)
        applicant_table.style = "Table Grid"
        
        # En-têtes
        applicant_cells = [
            ("Nom", applicant.get("name", "")),
            ("Adresse", applicant.get("address", "")),
            ("Téléphone", applicant.get("phone", "")),
            ("Email", applicant.get("email", "")),
            ("SIRET", applicant.get("siret", ""))
        ]
        
        for i, (header, value) in enumerate(applicant_cells):
            cell_header = applicant_table.cell(i, 0)
            cell_header.text = header
            cell_header.paragraphs[0].runs[0].bold = True
            
            cell_value = applicant_table.cell(i, 1)
            cell_value.text = str(value)
        
        doc.add_paragraph()  # Espace
        
        # Informations sur la parcelle
        doc.add_heading("Informations sur la parcelle", level=2)
        
        parcel_table = doc.add_table(rows=5, cols=2)
        parcel_table.style = "Table Grid"
        
        # En-têtes
        parcel_cells = [
            ("Identifiant", parcel.get("id", "")),
            ("Commune", parcel.get("commune", "")),
            ("Section", parcel.get("section", "")),
            ("Surface (ha)", f"{parcel.get('surface', 0):.2f}"),
            ("Région", parcel.get("region", ""))
        ]
        
        for i, (header, value) in enumerate(parcel_cells):
            cell_header = parcel_table.cell(i, 0)
            cell_header.text = header
            cell_header.paragraphs[0].runs[0].bold = True
            
            cell_value = parcel_table.cell(i, 1)
            cell_value.text = str(value)
        
        doc.add_paragraph()  # Espace
        
        # Informations sur le projet
        doc.add_heading("Informations sur le projet", level=2)
        
        project_table = doc.add_table(rows=5, cols=2)
        project_table.style = "Table Grid"
        
        # En-têtes
        project_cells = [
            ("Type de projet", project.get("type", "")),
            ("Budget total (€)", f"{project.get('budget', 0):.2f}"),
            ("Objectifs", project.get("objectives", "")),
            ("Date de début", project.get("start_date", "")),
            ("Durée (mois)", project.get("duration", ""))
        ]
        
        for i, (header, value) in enumerate(project_cells):
            cell_header = project_table.cell(i, 0)
            cell_header.text = header
            cell_header.paragraphs[0].runs[0].bold = True
            
            cell_value = project_table.cell(i, 1)
            cell_value.text = str(value)
        
        doc.add_paragraph()  # Espace
        
        # Informations sur la subvention
        doc.add_heading("Informations sur la subvention", level=2)
        
        # Description
        if subsidy.get("description"):
            description_para = doc.add_paragraph()
            description_para.add_run("Description: ").bold = True
            description_para.add_run(subsidy["description"])
            doc.add_paragraph()  # Espace
        
        # Tableau des informations de la subvention
        subsidy_table = doc.add_table(rows=6, cols=2)
        subsidy_table.style = "Table Grid"
        
        # En-têtes
        subsidy_cells = [
            ("Source", subsidy.get("source", "")),
            ("Région", subsidy.get("region", "")),
            ("Taux de financement", subsidy.get("financing_rate", "")),
            ("Montant min. (€)", str(subsidy.get("min_amount", "")) if subsidy.get("min_amount") else ""),
            ("Montant max. (€)", str(subsidy.get("max_amount", "")) if subsidy.get("max_amount") else ""),
            ("Date limite", subsidy.get("application_deadline", ""))
        ]
        
        for i, (header, value) in enumerate(subsidy_cells):
            cell_header = subsidy_table.cell(i, 0)
            cell_header.text = header
            cell_header.paragraphs[0].runs[0].bold = True
            
            cell_value = subsidy_table.cell(i, 1)
            cell_value.text = str(value)
        
        doc.add_paragraph()  # Espace
        
        # Critères d'éligibilité
        if subsidy.get("eligibility_criteria"):
            doc.add_heading("Critères d'éligibilité", level=2)
            
            for criterion in subsidy["eligibility_criteria"]:
                criterion_para = doc.add_paragraph(style="List Bullet")
                criterion_para.add_run(criterion)
            
            doc.add_paragraph()  # Espace
        
        # Contact
        if subsidy.get("contact"):
            contact_para = doc.add_paragraph()
            contact_para.add_run("Contact pour cette subvention: ").bold = True
            contact_para.add_run(subsidy["contact"])
        
        # URL pour plus d'informations
        if subsidy.get("url"):
            url_para = doc.add_paragraph()
            url_para.add_run("Pour plus d'informations: ").bold = True
            url_para.add_run(subsidy["url"])
        
        # Enregistrer le document
        doc.save(output_path)
        self.logger.info(f"Dossier de demande DOCX généré: {output_path}")
        
        return output_path
    
    def generate_eligibility_report(self, data: Dict[str, Any], output_path: str) -> str:
        """
        Génère un rapport d'éligibilité au format DOCX.
        
        Args:
            data: Données pour le document
            output_path: Chemin où enregistrer le document
        
        Returns:
            Chemin du fichier généré
        """
        self.logger.info("Génération d'un rapport d'éligibilité DOCX")
        
        # Vérifier si un template existe
        template_path = os.path.join(self.templates_dir, "eligibility_report_template.docx")
        if os.path.exists(template_path):
            doc = Document(template_path)
            self.logger.info(f"Utilisation du template: {template_path}")
        else:
            doc = Document()
            self.logger.warning("Template non trouvé, création d'un nouveau document")
        
        # Extraire les données nécessaires
        parcel = data["parcel"]
        project = data["project"]
        eligible_subsidies = data.get("eligible_subsidies", [])
        generated_at = data.get("generated_at", datetime.now().strftime("%d/%m/%Y"))
        
        # Formatage du document
        self._format_document_styles(doc)
        
        # Titre du document
        title = doc.add_heading("Rapport d'éligibilité aux subventions forestières", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date de génération
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(f"Document généré le {generated_at}")
        
        doc.add_paragraph()  # Espace
        
        # Informations sur la parcelle
        doc.add_heading("Informations sur la parcelle", level=2)
        
        parcel_table = doc.add_table(rows=5, cols=2)
        parcel_table.style = "Table Grid"
        
        # En-têtes
        parcel_cells = [
            ("Identifiant", parcel.get("id", "")),
            ("Commune", parcel.get("commune", "")),
            ("Section", parcel.get("section", "")),
            ("Surface (ha)", f"{parcel.get('surface', 0):.2f}"),
            ("Région", parcel.get("region", ""))
        ]
        
        for i, (header, value) in enumerate(parcel_cells):
            cell_header = parcel_table.cell(i, 0)
            cell_header.text = header
            cell_header.paragraphs[0].runs[0].bold = True
            
            cell_value = parcel_table.cell(i, 1)
            cell_value.text = str(value)
        
        doc.add_paragraph()  # Espace
        
        # Informations sur le projet
        doc.add_heading("Informations sur le projet", level=2)
        
        project_table = doc.add_table(rows=3, cols=2)
        project_table.style = "Table Grid"
        
        # En-têtes
        project_cells = [
            ("Type de projet", project.get("type", "")),
            ("Budget total (€)", f"{project.get('budget', 0):.2f}"),
            ("Objectifs", project.get("objectives", ""))
        ]
        
        for i, (header, value) in enumerate(project_cells):
            cell_header = project_table.cell(i, 0)
            cell_header.text = header
            cell_header.paragraphs[0].runs[0].bold = True
            
            cell_value = project_table.cell(i, 1)
            cell_value.text = str(value)
        
        doc.add_paragraph()  # Espace
        
        # Subventions éligibles
        doc.add_heading(f"Subventions éligibles ({len(eligible_subsidies)})", level=2)
        
        if eligible_subsidies:
            # Tableau récapitulatif
            summary_table = doc.add_table(rows=len(eligible_subsidies) + 1, cols=4)
            summary_table.style = "Table Grid"
            
            # En-tête du tableau
            headers = ["Subvention", "Source", "Score", "Montant estimé"]
            for i, header in enumerate(headers):
                cell = summary_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # Données du tableau
            for i, subsidy_result in enumerate(eligible_subsidies, 1):
                subsidy = subsidy_result["subsidy"]
                score = subsidy_result["eligibility_score"]
                amount = subsidy_result.get("estimated_amount", {}).get("estimated_amount", 0)
                
                summary_table.cell(i, 0).text = subsidy["title"]
                summary_table.cell(i, 1).text = subsidy.get("source", "")
                summary_table.cell(i, 2).text = f"{score:.2f}"
                summary_table.cell(i, 3).text = f"{amount:.2f} €"
            
            doc.add_paragraph()  # Espace
            
            # Détails des subventions
            doc.add_heading("Détails des subventions éligibles", level=2)
            
            for idx, subsidy_result in enumerate(eligible_subsidies, 1):
                subsidy = subsidy_result["subsidy"]
                score = subsidy_result["eligibility_score"]
                reasons = subsidy_result.get("reasons", "")
                requirements = subsidy_result.get("requirements", [])
                estimated = subsidy_result.get("estimated_amount", {})
                
                # Titre et détails de base
                doc.add_heading(f"{idx}. {subsidy['title']}", level=3)
                
                source_para = doc.add_paragraph()
                source_para.add_run("Source: ").bold = True
                source_para.add_run(subsidy.get("source", ""))
                
                score_para = doc.add_paragraph()
                score_para.add_run("Score d'éligibilité: ").bold = True
                score_para.add_run(f"{score:.2f}")
                
                # Description
                if subsidy.get("description"):
                    desc_para = doc.add_paragraph()
                    desc_para.add_run("Description: ").bold = True
                    desc_para.add_run(subsidy["description"])
                
                # Montant estimé
                if estimated:
                    amount_para = doc.add_paragraph()
                    amount_para.add_run("Montant estimé: ").bold = True
                    amount_para.add_run(f"{estimated.get('estimated_amount', 0):.2f} €")
                    
                    if "amount_explanation" in estimated:
                        expl_para = doc.add_paragraph()
                        expl_para.add_run("Explication: ").bold = True
                        expl_para.add_run(estimated["amount_explanation"])
                
                # Raisons d'éligibilité
                if reasons:
                    reasons_para = doc.add_paragraph()
                    reasons_para.add_run("Raisons d'éligibilité: ").bold = True
                    reasons_para.add_run(reasons)
                
                # Exigences
                if requirements:
                    doc.add_paragraph("Exigences à respecter:").paragraph_format.space_after = Pt(6)
                    
                    for req in requirements:
                        req_para = doc.add_paragraph(style="List Bullet")
                        req_para.add_run(req)
                
                # Contact
                if subsidy.get("contact"):
                    contact_para = doc.add_paragraph()
                    contact_para.add_run("Contact: ").bold = True
                    contact_para.add_run(subsidy["contact"])
                
                # URL
                if subsidy.get("url"):
                    url_para = doc.add_paragraph()
                    url_para.add_run("Plus d'informations: ").bold = True
                    url_para.add_run(subsidy["url"])
                
                # Séparateur
                if idx < len(eligible_subsidies):
                    doc.add_paragraph()
        else:
            doc.add_paragraph("Aucune subvention éligible trouvée pour ce projet.")
        
        # Enregistrer le document
        doc.save(output_path)
        self.logger.info(f"Rapport d'éligibilité DOCX généré: {output_path}")
        
        return output_path
    
    def generate_subsidies_summary(self, data: Dict[str, Any], output_path: str) -> str:
        """
        Génère une synthèse des subventions disponibles au format DOCX.
        
        Args:
            data: Données pour le document
            output_path: Chemin où enregistrer le document
        
        Returns:
            Chemin du fichier généré
        """
        self.logger.info("Génération d'une synthèse des subventions DOCX")
        
        # Vérifier si un template existe
        template_path = os.path.join(self.templates_dir, "subsidies_summary_template.docx")
        if os.path.exists(template_path):
            doc = Document(template_path)
            self.logger.info(f"Utilisation du template: {template_path}")
        else:
            doc = Document()
            self.logger.warning("Template non trouvé, création d'un nouveau document")
        
        # Extraire les données nécessaires
        subsidies = data.get("subsidies", [])
        filters = data.get("filters", {})
        generated_at = data.get("generated_at", datetime.now().strftime("%d/%m/%Y"))
        
        # Formatage du document
        self._format_document_styles(doc)
        
        # Titre du document
        title = doc.add_heading(f"Synthèse des subventions forestières ({len(subsidies)})", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date de génération
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(f"Document généré le {generated_at}")
        
        doc.add_paragraph()  # Espace
        
        # Filtres appliqués
        if filters:
            doc.add_heading("Filtres appliqués", level=2)
            
            if "region" in filters:
                filter_para = doc.add_paragraph(style="List Bullet")
                filter_para.add_run("Région: ").bold = True
                filter_para.add_run(filters["region"])
            
            if "project_type" in filters:
                filter_para = doc.add_paragraph(style="List Bullet")
                filter_para.add_run("Type de projet: ").bold = True
                filter_para.add_run(filters["project_type"])
            
            if "min_amount" in filters:
                filter_para = doc.add_paragraph(style="List Bullet")
                filter_para.add_run("Montant minimum: ").bold = True
                filter_para.add_run(f"{filters['min_amount']} €")
            
            if "keywords" in filters and filters["keywords"]:
                filter_para = doc.add_paragraph(style="List Bullet")
                filter_para.add_run("Mots-clés: ").bold = True
                filter_para.add_run(", ".join(filters["keywords"]))
            
            doc.add_paragraph()  # Espace
        
        # Tableau récapitulatif des subventions
        if subsidies:
            doc.add_heading("Récapitulatif des subventions", level=2)
            
            summary_table = doc.add_table(rows=len(subsidies) + 1, cols=5)
            summary_table.style = "Table Grid"
            
            # En-tête du tableau
            headers = ["Subvention", "Source", "Région", "Taux", "Montant max."]
            for i, header in enumerate(headers):
                cell = summary_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # Données du tableau
            for i, subsidy in enumerate(subsidies, 1):
                summary_table.cell(i, 0).text = subsidy["title"]
                summary_table.cell(i, 1).text = subsidy.get("source", "")
                summary_table.cell(i, 2).text = subsidy.get("region", "")
                summary_table.cell(i, 3).text = subsidy.get("financing_rate", "")
                summary_table.cell(i, 4).text = f"{subsidy.get('max_amount', '')} €" if subsidy.get("max_amount") else ""
            
            doc.add_paragraph()  # Espace
            
            # Détails des subventions
            doc.add_heading("Détails des subventions", level=2)
            
            for idx, subsidy in enumerate(subsidies, 1):
                # Titre et détails de base
                doc.add_heading(f"{idx}. {subsidy['title']}", level=3)
                
                source_para = doc.add_paragraph()
                source_para.add_run("Source: ").bold = True
                source_para.add_run(subsidy.get("source", ""))
                
                region_para = doc.add_paragraph()
                region_para.add_run("Région: ").bold = True
                region_para.add_run(subsidy.get("region", ""))
                
                # Description
                if subsidy.get("description"):
                    desc_para = doc.add_paragraph()
                    desc_para.add_run("Description: ").bold = True
                    desc_para.add_run(subsidy["description"])
                
                # Financement
                financing_info = []
                if subsidy.get("financing_rate"):
                    financing_info.append(f"Taux de financement: {subsidy['financing_rate']}")
                if subsidy.get("min_amount"):
                    financing_info.append(f"Montant minimum: {subsidy['min_amount']} €")
                if subsidy.get("max_amount"):
                    financing_info.append(f"Montant maximum: {subsidy['max_amount']} €")
                
                if financing_info:
                    doc.add_paragraph("Financement:").paragraph_format.space_after = Pt(6)
                    
                    for info in financing_info:
                        info_para = doc.add_paragraph(style="List Bullet")
                        info_para.add_run(info)
                
                # Projets éligibles
                if subsidy.get("eligible_projects"):
                    projects_para = doc.add_paragraph()
                    projects_para.add_run("Projets éligibles: ").bold = True
                    projects_para.add_run(", ".join(subsidy["eligible_projects"]))
                
                # Date limite
                if subsidy.get("deadline"):
                    deadline_para = doc.add_paragraph()
                    deadline_para.add_run("Date limite: ").bold = True
                    deadline_para.add_run(subsidy["deadline"])
                
                # Contact
                if subsidy.get("contact"):
                    contact_para = doc.add_paragraph()
                    contact_para.add_run("Contact: ").bold = True
                    contact_para.add_run(subsidy["contact"])
                
                # URL
                if subsidy.get("url"):
                    url_para = doc.add_paragraph()
                    url_para.add_run("Plus d'informations: ").bold = True
                    url_para.add_run(subsidy["url"])
                
                # Séparateur
                if idx < len(subsidies):
                    doc.add_paragraph()
        else:
            doc.add_paragraph("Aucune subvention disponible correspondant aux critères.")
        
        # Enregistrer le document
        doc.save(output_path)
        self.logger.info(f"Synthèse des subventions DOCX générée: {output_path}")
        
        return output_path
    
    def _format_document_styles(self, doc: Document) -> None:
        """
        Configure les styles du document.
        
        Args:
            doc: Document à formater
        """
        # Configuration des marges (si possible)
        try:
            for section in doc.sections:
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
        except:
            self.logger.warning("Impossible de configurer les marges du document")
        
        # Styles des paragraphes
        try:
            styles = doc.styles
            
            # Style Titre 1
            if "Heading 1" in styles:
                h1_style = styles["Heading 1"]
                h1_style.font.size = Pt(18)
                h1_style.font.bold = True
                h1_style.font.color.rgb = RGBColor(0, 51, 102)  # Bleu foncé
            
            # Style Titre 2
            if "Heading 2" in styles:
                h2_style = styles["Heading 2"]
                h2_style.font.size = Pt(16)
                h2_style.font.bold = True
                h2_style.font.color.rgb = RGBColor(0, 51, 102)  # Bleu foncé
            
            # Style Titre 3
            if "Heading 3" in styles:
                h3_style = styles["Heading 3"]
                h3_style.font.size = Pt(14)
                h3_style.font.bold = True
                h3_style.font.color.rgb = RGBColor(0, 102, 153)  # Bleu
            
            # Style Normal
            if "Normal" in styles:
                normal_style = styles["Normal"]
                normal_style.font.size = Pt(11)
                normal_style.font.name = "Calibri"
        except:
            self.logger.warning("Impossible de configurer les styles du document")
