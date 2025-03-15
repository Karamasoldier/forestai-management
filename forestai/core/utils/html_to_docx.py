# -*- coding: utf-8 -*-
"""
Module de conversion HTML vers DOCX.

Ce module fournit des fonctionnalités pour convertir du contenu HTML en DOCX,
en utilisant python-docx comme backend.
"""

import logging
import re
import io
import tempfile
import os
from pathlib import Path
from typing import Optional, Union, BytesIO, List

logger = logging.getLogger(__name__)

def html_to_docx(html_content: str, options: Optional[dict] = None) -> bytes:
    """
    Convertit du contenu HTML en document DOCX.
    
    Cette fonction utilise python-docx pour la génération de documents DOCX,
    avec une conversion simplifiée du HTML vers les structures DOCX.
    
    Args:
        html_content: Contenu HTML à convertir
        options: Options de conversion spécifiques
        
    Returns:
        Contenu DOCX en bytes
        
    Raises:
        RuntimeError: Si la conversion échoue
    """
    options = options or {}
    
    try:
        # On utilise une conversion simplifiée car une conversion HTML complète 
        # avec tous les styles est complexe
        return _simplified_html_to_docx(html_content, options)
    except Exception as e:
        logger.error(f"Erreur lors de la conversion HTML en DOCX: {str(e)}", exc_info=True)
        raise RuntimeError(f"Impossible de convertir le HTML en DOCX: {str(e)}")

def _simplified_html_to_docx(html_content: str, options: dict) -> bytes:
    """
    Effectue une conversion HTML vers DOCX simplifiée avec python-docx.
    
    Args:
        html_content: Contenu HTML à convertir
        options: Options de conversion
        
    Returns:
        Contenu DOCX en bytes
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from bs4 import BeautifulSoup
        
        # Créer un document vide
        document = Document()
        
        # Définir les marges du document
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Parser le HTML avec BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extraire le titre depuis une balise title ou h1
        title_tag = soup.find('title')
        if not title_tag:
            title_tag = soup.find('h1')
        
        if title_tag:
            # Ajouter le titre au document
            title = document.add_heading(title_tag.get_text().strip(), level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fonction récursive pour traiter les éléments HTML
        def process_element(element, parent_docx_element=document):
            if element.name is None:  # Text node
                if element.strip():
                    if parent_docx_element == document:
                        p = parent_docx_element.add_paragraph()
                        p.add_run(element.strip())
                    else:
                        parent_docx_element.add_run(element.strip())
            
            elif element.name == 'p':
                p = document.add_paragraph()
                for child in element.children:
                    process_element(child, p)
            
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                heading = document.add_heading(element.get_text().strip(), level=level)
            
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    p = document.add_paragraph(style='List Bullet')
                    for child in li.children:
                        if child.name != 'ul' and child.name != 'ol':  # Éviter les listes imbriquées
                            process_element(child, p)
            
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    p = document.add_paragraph(style='List Number')
                    for child in li.children:
                        if child.name != 'ul' and child.name != 'ol':  # Éviter les listes imbriquées
                            process_element(child, p)
            
            elif element.name == 'a':
                # Ajouter un lien hypertexte
                if parent_docx_element == document:
                    p = parent_docx_element.add_paragraph()
                    p.add_run(element.get_text().strip()).underline = True
                else:
                    run = parent_docx_element.add_run(element.get_text().strip())
                    run.underline = True
            
            elif element.name == 'strong' or element.name == 'b':
                # Texte en gras
                if parent_docx_element == document:
                    p = parent_docx_element.add_paragraph()
                    p.add_run(element.get_text().strip()).bold = True
                else:
                    parent_docx_element.add_run(element.get_text().strip()).bold = True
            
            elif element.name == 'em' or element.name == 'i':
                # Texte en italique
                if parent_docx_element == document:
                    p = parent_docx_element.add_paragraph()
                    p.add_run(element.get_text().strip()).italic = True
                else:
                    parent_docx_element.add_run(element.get_text().strip()).italic = True
            
            elif element.name == 'table':
                # Tableau simple
                rows = element.find_all('tr')
                if rows:
                    table = document.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                    table.style = 'Table Grid'
                    
                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < len(table.rows[i].cells):  # Vérifier les limites
                                table.rows[i].cells[j].text = cell.get_text().strip()
            
            elif element.name == 'br':
                # Saut de ligne
                if parent_docx_element != document:
                    parent_docx_element.add_run('\n')
            
            elif element.name == 'div':
                # Traiter les div comme des conteneurs
                for child in element.children:
                    process_element(child, document)
            
            elif element.name == 'hr':
                # Ligne horizontale
                document.add_paragraph('_' * 50)
            
            else:
                # Traiter les autres éléments comme du texte simple
                for child in element.children:
                    process_element(child, parent_docx_element)
        
        # Traiter le corps du document (uniquement les éléments du body)
        body = soup.find('body')
        if body:
            for child in body.children:
                process_element(child)
        else:
            # Si pas de body, traiter tous les éléments de premier niveau
            for child in soup.children:
                if child.name != 'head' and child.name != 'title':
                    process_element(child)
        
        # Sauvegarder le document dans un buffer
        docx_buffer = BytesIO()
        document.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer.getvalue()
    
    except ImportError:
        raise RuntimeError("python-docx ou BeautifulSoup n'est pas installé.")

def _mammoth_html_to_docx(html_content: str, options: dict) -> bytes:
    """
    Utilise mammoth (inversé) pour convertir HTML en DOCX.
    Cette approche est une alternative mais peut nécessiter des bibliothèques supplémentaires.
    
    Args:
        html_content: Contenu HTML
        options: Options de conversion
        
    Returns:
        Contenu DOCX en bytes
    """
    try:
        import mammoth
        from tempfile import NamedTemporaryFile
        
        # Écrire le HTML dans un fichier temporaire
        with NamedTemporaryFile(mode='w', suffix='.html', delete=False) as temp_html:
            temp_html.write(html_content)
            temp_html_path = temp_html.name
        
        # Créer un fichier temporaire pour le DOCX
        with NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx_path = temp_docx.name
        
        try:
            # Convertir HTML en DOCX (cette fonctionnalité n'est pas standard dans mammoth,
            # cette partie est hypothétique et nécessiterait une implémentation personnalisée)
            # mammoth.convert_html_to_docx(temp_html_path, temp_docx_path)
            
            # Lire le fichier DOCX généré
            with open(temp_docx_path, 'rb') as f:
                docx_content = f.read()
            
            return docx_content
        
        finally:
            # Nettoyer les fichiers temporaires
            if os.path.exists(temp_html_path):
                os.unlink(temp_html_path)
            if os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
    
    except ImportError:
        raise RuntimeError("mammoth n'est pas installé.")
